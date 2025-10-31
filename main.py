import sys, os, subprocess, json, uuid, requests, re
from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional
from decimal import Decimal, getcontext
from docx import Document
from datetime import datetime
from utils.parser import extraer_variables

# -------------------------------------------------
# CONFIGURACIÓN INICIAL
# -------------------------------------------------
getcontext().prec = 28

app = FastAPI(title="TemplatePDF - Cotización y generación de PDF")

TEMPLATES_DIR = "templates"
OUTPUT_DIR = "outputs"
DB_PATH = "db.json"
TEMPLATE_NAME = "Plantilla_Cotizacion.docx"
GITHUB_RAW_URL = "https://github.com/FirstLeaseAgent/template-pdf/raw/refs/heads/main/templates/Plantilla_Cotizacion.docx"

os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Crear base de datos si no existe
if not os.path.exists(DB_PATH):
    with open(DB_PATH, "w") as f:
        json.dump({"plantillas": []}, f, indent=4)


# -------------------------------------------------
# AUTO-CARGA DE PLANTILLA DESDE GITHUB
# -------------------------------------------------
def ensure_template_available():
    template_path = os.path.join(TEMPLATES_DIR, TEMPLATE_NAME)
    if not os.path.exists(template_path):
        print("🔄 Descargando plantilla desde GitHub...")
        resp = requests.get(GITHUB_RAW_URL)
        resp.raise_for_status()
        with open(template_path, "wb") as f:
            f.write(resp.content)
        print("✅ Plantilla descargada correctamente.")

    with open(DB_PATH, "r+") as db_file:
        data = json.load(db_file)
        if not data["plantillas"]:
            plantilla_id = str(uuid.uuid4())
            data["plantillas"].append({
                "id": plantilla_id,
                "nombre": TEMPLATE_NAME,
                "variables": []
            })
            db_file.seek(0)
            db_file.truncate()
            json.dump(data, db_file, indent=4)
            print("✅ Registro de plantilla agregado a db.json")

ensure_template_available()



# -------------------------------------------------
# FUNCIONES DE CÁLCULO Y UTILIDADES
# -------------------------------------------------
class Activo(BaseModel):
    nombre_activo: str
    valor: float
    enganche: Optional[float] = 10.0
    tasa_anual: Optional[float] = 30.0
    comision: Optional[float] = 3.0
    rentas_deposito: Optional[float] = 1.0

class CotizacionRequest(BaseModel):
    nombre: str
    activos: List[Activo]


def formato_miles(valor):
    try:
        num = float(valor)
        return f"{num:,.2f}"
    except:
        return valor


def calcular_pago_mensual(valor, enganche, tasa_anual, plazo_meses, valor_residual, comision, rentas_deposito):
    pv = Decimal(valor / 1.16) * Decimal(1 - enganche / 100)
    r = Decimal(tasa_anual) / Decimal(100 * 12)
    n = Decimal(plazo_meses)
    fv = Decimal(valor / 1.16 * valor_residual / 100)

    if r == 0:
        pago = -(pv - fv) / n
    else:
        pago = ((pv - fv * ((1 + r) ** (-n))) * r) / (1 - (1 + r) ** (-n))

    monto_comision = Decimal(comision) / Decimal(100) * pv
    monto_enganche = Decimal(enganche) / Decimal(100) * Decimal(valor) / Decimal("1.16")
    monto_deposito = Decimal(rentas_deposito) * pago * Decimal("1.16")
    monto_residual = Decimal(valor) / Decimal("1.16") * Decimal(valor_residual) / Decimal(100)

    subtotal_inicial = monto_enganche + monto_comision + monto_deposito + pago
    iva_inicial = (monto_enganche + monto_comision + pago) * Decimal("0.16")
    total_inicial = subtotal_inicial + iva_inicial

    iva_renta = pago * Decimal("0.16")
    total_renta = pago * Decimal("1.16")

    iva_residual = monto_residual * Decimal("0.16")
    total_residual = monto_residual * Decimal("1.16")

    total_final = total_residual - monto_deposito

    return {
        "Enganche": float(round(monto_enganche, 2)),
        "Comision": float(round(monto_comision, 2)),
        "Renta_en_Deposito": float(round(monto_deposito, 2)),
        "Subtotal_Pago_Inicial": float(round(subtotal_inicial, 2)),
        "IVA_Pago_Inicial": float(round(iva_inicial, 2)),
        "Total_Inicial": float(round(total_inicial, 2)),
        "Renta_Mensual": float(round(pago, 2)),
        "IVA_Renta_Mensual": float(round(iva_renta, 2)),
        "Total_Renta_Mensual": float(round(total_renta, 2)),
        "Residual": float(round(monto_residual, 2)),
        "IVA_Residual": float(round(iva_residual, 2)),
        "Total_Residual": float(round(total_residual, 2)),
        "Reembolso_Deposito": float(round(-monto_deposito, 2)),
        "Total_Final": float(round(total_final, 2))
    }


# -------------------------------------------------
# FUNCIÓN: CONVERTIR WORD A PDF (LIBREOFFICE)
# -------------------------------------------------
def convertir_a_pdf(word_path, output_dir):
    """
    Convierte un archivo .docx a PDF usando LibreOffice.
    """
    try:
        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            word_path
        ], check=True)
        pdf_name = os.path.splitext(os.path.basename(word_path))[0] + ".pdf"
        return os.path.join(output_dir, pdf_name)
    except subprocess.CalledProcessError as e:
        raise Exception(f"Error al convertir a PDF: {e}")


# -------------------------------------------------
# ENDPOINT PRINCIPAL /cotizar
# -------------------------------------------------
@app.post("/cotizar")
def cotizar(data: CotizacionRequest, request: Request):
    escenarios = [
        {"plazo": 24, "residual": 40},
        {"plazo": 36, "residual": 30},
        {"plazo": 48, "residual": 25},
    ]

    resultado = {"Nombre": data.nombre, "Cotizaciones": []}

    valores_para_doc = {
        "nombre": data.nombre,
        "descripcion": "",
        "precio": "",
        "fecha": datetime.now().strftime("%d/%m/%Y"),
        "folio": datetime.now().strftime("%Y%m%d%H%M%S"),
    }

    for activo in data.activos:
        valores_para_doc["descripcion"] = activo.nombre_activo
        valores_para_doc["precio"] = formato_miles(activo.valor)

        for e in escenarios:
            calculos = calcular_pago_mensual(
                valor=activo.valor,
                enganche=activo.enganche,
                tasa_anual=activo.tasa_anual,
                plazo_meses=e["plazo"],
                valor_residual=e["residual"],
                comision=activo.comision,
                rentas_deposito=activo.rentas_deposito,
            )

            # 🧩 Mapeo de alias para nombres esperados en la plantilla
            alias = {
                "rentaendeposito": "deposito",
                "rentamensual": "mensualidad",
                "ivarentamensual": "IVAmes",
                "subtotalpagoinicial": "subinicial",
                "ivapagoinicial": "IVAinicial",
                "ivaresidual": "IVAresidual",
                "reembolsodeposito": "reembolso",
                "totalrentamensual": "totalmes"
            }

            # 🧮 Generar variables para la plantilla (24/36/48)
            plazo = str(e["plazo"])
            for k, v in calculos.items():
                clave = k.replace("_", "").lower()
                nombre_final = alias.get(clave, clave)
                valores_para_doc[f"{nombre_final}{plazo}"] = formato_miles(v)

            # 📊 Guardar también los cálculos para el JSON de salida
            resultado["Cotizaciones"].append(calculos)

    # ==============================
    # Generar documento Word + PDF
    # ==============================

    # Buscar plantilla registrada
    with open(DB_PATH, "r") as db_file:
        db_data = json.load(db_file)

    plantilla = db_data["plantillas"][0] if db_data["plantillas"] else None

    if not plantilla:
        raise HTTPException(status_code=404, detail="No hay plantilla registrada")

    # Generar documento y PDF
    documentos = generar_documento_word_local(
        plantilla_id=plantilla["id"],
        valores=valores_para_doc,
        request=request
    )

    return {
        "mensaje": "Cotización generada correctamente",
        "folio": valores_para_doc["folio"],
        "cotizaciones": resultado["Cotizaciones"],
        "documentos": documentos
    }

# -------------------------------------------------
# DEBUG: detectar marcadores en Word
# -------------------------------------------------
def debug_list_placeholders(doc_path):
    """
    Escanea un archivo .docx y lista todos los placeholders {{...}} detectados,
    incluso si están divididos en runs.
    """

    print(f"\n🔍 Analizando marcadores en: {doc_path}")
    doc = Document(doc_path)

    pattern = re.compile(r"\{\{(.*?)\}\}")
    encontrados = set()

    # Párrafos normales
    for p in doc.paragraphs:
        text = "".join(run.text for run in p.runs)
        matches = pattern.findall(text)
        for m in matches:
            encontrados.add(m.strip())

    # Celdas de tabla
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = "".join(run.text for p in cell.paragraphs for run in p.runs)
                matches = pattern.findall(text)
                for m in matches:
                    encontrados.add(m.strip())

    print("📄 Marcadores detectados en la plantilla:")
    for m in sorted(encontrados):
        print("   •", m)

    print(f"🧾 Total: {len(encontrados)} marcadores encontrados.\n")

    ensure_template_available()
    debug_list_placeholders(os.path.join(TEMPLATES_DIR, TEMPLATE_NAME))
    
    # --- Generar documento Word ---
def generar_documento_word_local(plantilla_id: str, valores: dict, request: Request):


    # 1. Cargar DB
    with open(DB_PATH, "r") as f:
        data = json.load(f)

    # 2. Buscar plantilla por ID
    plantilla = next((p for p in data["plantillas"] if p["id"] == plantilla_id), None)
    if not plantilla:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")

    plantilla_path = os.path.join(TEMPLATES_DIR, plantilla["nombre"])

    # 3. Si la plantilla no existe localmente, descargarla desde GitHub
    if not os.path.exists(plantilla_path):
        GITHUB_RAW_URL = "https://github.com/FirstLeaseAgent/template-pdf/raw/refs/heads/main/templates/Plantilla_Cotizacion.docx"
        try:
            response = requests.get(GITHUB_RAW_URL, timeout=30)
            response.raise_for_status()
            with open(plantilla_path, "wb") as f:
                f.write(response.content)
            print(f"✅ Plantilla descargada desde GitHub: {plantilla['nombre']}")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error al descargar plantilla desde GitHub: {e}")

    # 4. Cargar Word
    doc = Document(plantilla_path)
    debug_list_placeholders(plantilla_path)
    # 5. Reemplazo de variables manteniendo formato
    for p in doc.paragraphs:
        for run in p.runs:
            for var, valor in valores.items():
                placeholder = f"{{{{{var}}}}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(valor))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for var, valor in valores.items():
                            placeholder = f"{{{{{var}}}}}"
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(valor))

    # 6. Guardar archivo Word
    folio = valores.get("folio", datetime.now().strftime("%Y%m%d_%H%M%S"))
    word_name = f"cotizacion_{folio}.docx"
    word_path = os.path.join(OUTPUT_DIR, word_name)
    doc.save(word_path)

    # 7. Convertir a PDF con LibreOffice
    pdf_name = word_name.replace(".docx", ".pdf")
    pdf_path = os.path.join(OUTPUT_DIR, pdf_name)
    try:
        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", OUTPUT_DIR,
            word_path
        ], check=True)
        print(f"✅ PDF generado correctamente: {pdf_name}")
    except Exception as e:
        print(f"⚠️ Error al generar PDF: {e}")
        pdf_path = None

    # 8. Construir URLs de descarga
    base_url = str(request.base_url).rstrip("/")
    download_word = f"{base_url}/download_word/{word_name}"
    download_pdf = f"{base_url}/download_word/{pdf_name}" if pdf_path else None

    print(f"🧾 Documento generado con folio {folio}")

    return {
        "archivo_word": word_name,
        "descargar_word": download_word,
        "archivo_pdf": pdf_name if pdf_path else "No generado",
        "descargar_pdf": download_pdf,
        "folio": folio
    }

# -------------------------------------------------
# ENDPOINTS DE DESCARGA
# -------------------------------------------------
@app.get("/download_word/{filename}")
def download_word(filename: str):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    return FileResponse(path, filename=filename)


@app.get("/download_pdf/{filename}")
def download_pdf(filename: str):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    return FileResponse(path, filename=filename)


@app.get("/")
def root():
    return {"mensaje": "TemplatePDF API funcionando correctamente 🚀"}

# -------------------------------------------------
# ENDPOINT PARA RECARGAR PLANTILLA DESDE GITHUB
# -------------------------------------------------

@app.get("/reload_template")
def reload_template():
    """
    Fuerza la descarga de la plantilla desde GitHub y reemplaza la versión local.
    """
    import requests

    template_path = os.path.join(TEMPLATES_DIR, TEMPLATE_NAME)
    try:
        # Descargar la plantilla actualizada desde GitHub
        resp = requests.get(GITHUB_RAW_URL, timeout=30)
        resp.raise_for_status()

        # Guardar la nueva versión localmente
        with open(template_path, "wb") as f:
            f.write(resp.content)

        # Actualizar el registro en db.json
        with open(DB_PATH, "r+") as db_file:
            data = json.load(db_file)
            if data["plantillas"]:
                data["plantillas"][0]["nombre"] = TEMPLATE_NAME
            else:
                data["plantillas"].append({
                    "id": str(uuid.uuid4()),
                    "nombre": TEMPLATE_NAME,
                    "variables": []
                })
            db_file.seek(0)
            db_file.truncate()
            json.dump(data, db_file, indent=4)

        return {"mensaje": "✅ Plantilla actualizada correctamente desde GitHub"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al actualizar plantilla: {str(e)}")