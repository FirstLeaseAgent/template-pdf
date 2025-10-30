import re
from docx import Document

def extraer_variables(file_path):
    """
    Extrae variables {{variable}} de una plantilla Word (.docx)
    """
    doc = Document(file_path)
    texto = ""
    for p in doc.paragraphs:
        texto += p.text + " "
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto += cell.text + " "

    # Buscar variables con doble llave {{var}}
    variables = re.findall(r"{{(.*?)}}", texto)
    return list(set(variables))